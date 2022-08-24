from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    """Метод, добавляющий номера страниц в word."""
    # выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fld_char1 = create_element('w:fldChar')
    create_attribute(fld_char1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUM PAGES)
    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "PAGE"
    # обозначаем конец позиции вывода
    fld_char2 = create_element('w:fldChar')
    create_attribute(fld_char2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fld_char1)
    page_num_run._r.append(instr_text)
    page_num_run._r.append(fld_char2)


def add_new_styles(document):
    """Метод, добавляющий стили текста."""
    styles = document.styles
    styles.add_style('PItalic', WD_STYLE_TYPE.PARAGRAPH)
    style = document.styles['PItalic']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.italic = True

    styles.add_style('Head-style', WD_STYLE_TYPE.PARAGRAPH)
    head_style = document.styles['Head-style']
    font = head_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True

    styles.add_style('Head-graphic', WD_STYLE_TYPE.PARAGRAPH)
    head_graphic = document.styles['Head-graphic']
    font = head_graphic.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.bold = True
    font.italic = True


def section_choice(document):
    """Метод, добавляющий отступы в документе word."""
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def make_table_bold(table, cols, rows):
    """Метод, изменяющий вес шрифтов в таблицах и выравнивающий таблицу по центру."""
    for row in range(1):
        for col in range(cols):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, rows):
        for col in range(1):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def make_stat_table_bold(stat_table, cols, rows):
    """Метод, изменяющий вес шрифтов в статистических таблицах и выравнивающий таблицу по центру."""
    for row in stat_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for row in range(2):
        for col in range(cols):
            # получаем ячейку таблицы
            cell = stat_table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

    for row in range(2, rows):
        for col in range(2):
            # получаем ячейку таблицы
            cell = stat_table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(8.5)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER


def draw_stat_table(stat_table, stat_info_1, stat_info_2):
    """Метод для построения статистических таблиц с информацией о средних значениях и о стандартных отклонениях
    параметра."""
    stat_table.cell(0, 0).text = "№"
    stat_table.cell(0, 1).text = "Стат-ка"
    stat_table.cell(0, 2).text = "№ детектора"
    stat_table.cell(0, 2).merge(stat_table.cell(0, 17))
    stat_table.cell(0, 0).merge(stat_table.cell(1, 0))
    stat_table.cell(0, 1).merge(stat_table.cell(1, 1))
    for j in range(stat_info_1.shape[-1]):
        stat_table.cell(1, j + 2).text = stat_info_1.columns[j]
    stat_table.cell(2, 0).text = '1'
    stat_table.cell(2, 0).merge(stat_table.cell(3, 0))
    stat_table.cell(4, 0).text = '2'
    stat_table.cell(4, 0).merge(stat_table.cell(5, 0))
    for i in range(stat_info_1.shape[0]):
        stat_table.cell(i + 2, 1).text = stat_info_1.index[i]
        for j in range(stat_info_1.shape[-1]):
            stat_table.cell(i + 2, j + 2).text = str(round(stat_info_1.values[i, j], 2))

    for i in range(stat_info_2.shape[0]):
        stat_table.cell(i + 2 + stat_info_1.shape[0], 1).text = stat_info_2.index[i]
        for j in range(stat_info_2.shape[-1]):
            stat_table.cell(i + 2 + stat_info_1.shape[0], j + 2).text = str(round(stat_info_2.values[i, j], 2))


def change_cell_size(table, column_num, size_arr):
    """Метод, меняющий размер клеток в таблице."""
    for i in range(column_num):
        for cell in table.columns[i].cells:
            cell.width = Inches(size_arr[i])


def adding_graphic(document, title, picture_path, width):
    """Метод, добавляющий в word график."""
    document.add_picture(picture_path, width=Inches(width))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_title = document.add_paragraph(title, style='PItalic')
    picture_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def statistical_table(frame, frame_2, dimension):
    """Метод, на вход которого, идут два датафрейма 1-го и 2-го кластера ПРИЗМА-32 с одинаковым параметром. Метод
    возвращает датафрейм с информацией о средних значениях и о стандартных отклонениях заданного параметра"""
    stat_info = frame.describe().tail(7).head(2)
    stat_info_2 = frame_2.describe().tail(7).head(2)
    stat_info.index = [f'mean({dimension})', f'std({dimension})']
    stat_info_2.index = [f'mean({dimension})', f'std({dimension})']
    return stat_info, stat_info_2


def time_breaks_counter(brake_frame):
    breaks = len(brake_frame.index)
    fail_str_begin = []
    fail_str_end = []
    lost_minutes = []
    for i in range(len(brake_frame.index)):
        fail_str_begin.append(
            f" {brake_frame['Date'][i].date()}  {brake_frame['StartMinutes'][i] // 60:02}:{brake_frame['StartMinutes'][i] % 60:02}")
        fail_str_end.append(
            f" {brake_frame['Date'][i].date()}  {brake_frame['EndMinutes'][i] // 60:02}:{brake_frame['EndMinutes'][i] % 60:02}")
        lost_minutes.append(brake_frame['EndMinutes'][i] - brake_frame['StartMinutes'][i])

    for i in range(1, len(brake_frame.index)):
        if brake_frame['StartMinutes'][i] == 0 and brake_frame['EndMinutes'][i - 1] == 1435 and \
                (brake_frame['Date'][i] - brake_frame['Date'][i-1]).days == 1:
            breaks -= 1

    return fail_str_begin, fail_str_end, lost_minutes, breaks
