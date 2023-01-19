import time
import warnings

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK

from drawing_graphs import GraphsDrawing
from processing_data_prisma import ProccessingPrismaCl
from word_addition import *

amp_n_cols = []
for item in range(1, 17):
    amp_n_cols.append(f'amp{item}')
    amp_n_cols.append(f'n{item}')


def make_report_prisma(start_date, end_date, report_path, picture_path, concat_n_df_1, concat_n_df_2):
    t1 = time.time()

    warnings.filterwarnings(action='ignore')

    days_amount = len(pd.date_range(start_date, end_date))

    process_1 = ProccessingPrismaCl(n_data=concat_n_df_1)
    process_2 = ProccessingPrismaCl(n_data=concat_n_df_2)

    graphs = GraphsDrawing(start_date=start_date, end_date=end_date,
                           path_to_pic=f'{picture_path}')

    worktime_frame_1, breaks_frame_1, n_vs_zero_tr_frame_1, event_counter_fr_4_1, count_rate_amp_5_fr_2_1,\
    count_rate_amp_10_fr_1_1, amp_5_fr_2_frame_1, amp_10_fr_1_frame_1, = process_1.period_processing_for_report(
        start_date=start_date, end_date=end_date)

    worktime_frame_2, breaks_frame_2, n_vs_zero_tr_frame_2, event_counter_fr_4_2, count_rate_amp_5_fr_2_2, \
    count_rate_amp_10_fr_1_2, amp_5_fr_2_frame_2, amp_10_fr_1_frame_2 = process_2.period_processing_for_report(
        start_date=start_date, end_date=end_date)

    brake_both_cl_time = 0
    for i in range(len(breaks_frame_1.index)):
        for j in range(len(breaks_frame_2.index)):
            if breaks_frame_1['Date'][i] == breaks_frame_2['Date'][j]:
                if breaks_frame_1['StartSeconds'][i] <= breaks_frame_2['EndSeconds'][j] < \
                        breaks_frame_1['EndSeconds'][i]:
                    brake_both_cl_time += min(breaks_frame_2['EndSeconds'][j], breaks_frame_1['EndSeconds'][i]) - max(
                        breaks_frame_2['StartSeconds'][j], breaks_frame_1['StartSeconds'][i])
                elif breaks_frame_2['StartSeconds'][j] <= breaks_frame_1['StartSeconds'][i] < \
                        breaks_frame_2['EndSeconds'][
                            j]:
                    brake_both_cl_time += min(breaks_frame_2['EndSeconds'][j], breaks_frame_1['EndSeconds'][i]) - max(
                        breaks_frame_2['StartSeconds'][j], breaks_frame_1['StartSeconds'][i])

    real_worktime = worktime_frame_2['Worktime'].sum() - 24 * days_amount + worktime_frame_1[
        'Worktime'].sum() + brake_both_cl_time / 3600

    print(f'{brake_both_cl_time=}')

    graphs.change_design()

    worktime_pic_path_1 = graphs.worktime_graph(cluster=1, worktime_frame=worktime_frame_1)
    worktime_pic_path_2 = graphs.worktime_graph(cluster=2, worktime_frame=worktime_frame_2)

    n_vs_zero_tr_pic_path_1 = graphs.neutron_to_0_tr_graph(cluster=1, neutron_num_0_tr_frame=n_vs_zero_tr_frame_1)
    n_vs_zero_tr_pic_path_2 = graphs.neutron_to_0_tr_graph(cluster=2, neutron_num_0_tr_frame=n_vs_zero_tr_frame_2)

    event_counter_fr_4_pic_path = graphs.amp_5_fr_4_graph(amp_5_fr_4_frame=event_counter_fr_4_1,
                                                          amp_5_fr_4_frame_2=event_counter_fr_4_2,
                                                          worktime_frame=worktime_frame_1,
                                                          worktime_frame_2=worktime_frame_2)

    amp_distribution_pic_path_1 = graphs.amp_distribution_graph(cluster=1, amp_distribution_frame=amp_5_fr_2_frame_1,
                                                                a_crit=6, freq=2)
    amp_distribution_pic_path_2 = graphs.amp_distribution_graph(cluster=2, amp_distribution_frame=amp_5_fr_2_frame_2,
                                                                a_crit=6, freq=2)

    count_rate_amp_5_fr_2_pic_path_1 = graphs.count_rate_graph(cluster=1, count_rate_frame=count_rate_amp_5_fr_2_1,
                                                               working_frame=worktime_frame_1,
                                                               a_crit=5, freq=2)
    count_rate_amp_5_fr_2_pic_path_2 = graphs.count_rate_graph(cluster=2, count_rate_frame=count_rate_amp_5_fr_2_2,
                                                               working_frame=worktime_frame_2,
                                                               a_crit=5, freq=2)

    count_rate_amp_10_fr_1_pic_path_1 = graphs.count_rate_graph(cluster=1, count_rate_frame=count_rate_amp_10_fr_1_1,
                                                                working_frame=worktime_frame_1,
                                                                a_crit=10, freq=1)
    count_rate_amp_10_fr_1_pic_path_2 = graphs.count_rate_graph(cluster=2, count_rate_frame=count_rate_amp_10_fr_1_2,
                                                                working_frame=worktime_frame_2,
                                                                a_crit=10, freq=1)

    del graphs

    doc = Document()
    section_choice(doc)
    add_new_styles(doc)

    head = doc.add_paragraph(
        f'Справка о работе установки ПРИЗМА-32 в период с {start_date} по {end_date} ', style='Head-style')
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table_title = doc.add_paragraph('Таблица 1: Время работы установки ПРИЗМА-32.', style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    worktime_table = doc.add_table(4, 4, doc.styles['Table Grid'])
    worktime_table.cell(0, 0).text = '№ кластера'
    worktime_table.cell(0, 1).text = 'Экспозиции, ч.'
    worktime_table.cell(0, 2).text = 'Календарное время, ч.'
    worktime_table.cell(0, 3).text = 'Экспозиция, %'
    worktime_table.cell(1, 0).text = '1'
    worktime_table.cell(1, 1).text = str(round(worktime_frame_1['Worktime'].sum(), 2))
    worktime_table.cell(1, 2).text = str(24 * days_amount)
    worktime_table.cell(1, 3).text = str(
        round(worktime_frame_1['Worktime'].sum() / (24 * days_amount) * 100, 3)) + '%'

    worktime_table.cell(2, 0).text = '2'
    worktime_table.cell(2, 1).text = str(round(worktime_frame_2['Worktime'].sum(), 2))
    worktime_table.cell(2, 2).text = str(24 * days_amount)
    worktime_table.cell(2, 3).text = str(
        round(worktime_frame_2['Worktime'].sum() / (24 * days_amount) * 100, 3)) + '%'

    worktime_table.cell(3, 0).text = '1&2'
    worktime_table.cell(3, 1).text = str(round(real_worktime, 2))
    worktime_table.cell(3, 2).text = str(24 * days_amount)
    worktime_table.cell(3, 3).text = str(round(real_worktime / (24 * days_amount) * 100, 3)) + '%'

    make_table_bold(worktime_table, cols=4, rows=4)

    doc.add_paragraph()

    fail_str_begin_1, fail_str_end_1, lost_minutes_1, break_1 = time_breaks_counter(brake_frame=breaks_frame_1)
    fail_str_begin_2, fail_str_end_2, lost_minutes_2, break_2 = time_breaks_counter(brake_frame=breaks_frame_2)
    brake_table_title = doc.add_paragraph('Таблица 2: Сводная таблица остановок и работ установки ПРИЗМА-32.',
                                          style='PItalic')
    brake_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    brake_table = doc.add_table(len(fail_str_begin_1) + len(fail_str_begin_2) + 2, 5, doc.styles['Table Grid'])
    brake_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    brake_table.cell(0, 0).text = '№ кластера'
    brake_table.cell(0, 0).merge(brake_table.cell(1, 0))
    brake_table.cell(0, 1).text = 'Время простоя'
    brake_table.cell(1, 1).text = 'c'
    brake_table.cell(1, 2).text = 'по'
    brake_table.cell(0, 1).merge(brake_table.cell(0, 2))
    brake_table.cell(0, 3).text = 'Кол-во потерянных минут (период)'
    brake_table.cell(0, 3).merge(brake_table.cell(1, 3))
    brake_table.cell(0, 4).text = 'Примечание'
    brake_table.cell(0, 4).merge(brake_table.cell(1, 4))

    for i in range(2, len(fail_str_begin_1) + 2):
        brake_table.cell(i, 0).text = '1'
        brake_table.cell(i, 1).text = str(fail_str_begin_1[i - 2])
        brake_table.cell(i, 2).text = str(fail_str_end_1[i - 2])
        brake_table.cell(i, 3).text = str(lost_minutes_1[i - 2])
        brake_table.cell(i, 4).text = ' '

    for i in range(2 + len(fail_str_begin_1), len(fail_str_begin_2) + 2 + len(fail_str_begin_1)):
        brake_table.cell(i, 0).text = '2'
        brake_table.cell(i, 1).text = str(fail_str_begin_2[i - 2 - len(fail_str_begin_1)])
        brake_table.cell(i, 2).text = str(fail_str_end_2[i - 2 - len(fail_str_begin_1)])
        brake_table.cell(i, 3).text = str(lost_minutes_2[i - 2 - len(fail_str_begin_1)])
        brake_table.cell(i, 4).text = ' '

    make_table_bold(brake_table, cols=5, rows=len(fail_str_begin_1) + len(fail_str_begin_2) + 2)
    doc.add_paragraph()

    table_title = doc.add_paragraph(
        'Таблица 3: Сводная таблица темпов счета событий и сигналов, отобранных как нейтрон кластеров установки ПРИЗМА-32.',
        style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    neut_stat_info_1, neut_stat_info_2 = statistical_table(n_vs_zero_tr_frame_1, n_vs_zero_tr_frame_2,
                                                           dimension='100/соб')

    neutron_table = doc.add_table(3, 3, doc.styles['Table Grid'])
    neutron_table.cell(0, 0).text = 'Счет/кластер'
    neutron_table.cell(0, 1).text = 'Кл1'
    neutron_table.cell(0, 2).text = 'Кл2'
    neutron_table.cell(1, 0).text = 'События (Fr ≥ 4, A ≥ 5), N соб./ч.'
    neutron_table.cell(1, 1).text = str(
        round((event_counter_fr_4_1['Events'] / worktime_frame_1['Worktime']).mean(), 2))
    neutron_table.cell(1, 2).text = str(
        round((event_counter_fr_4_2['Events'] / worktime_frame_2['Worktime']).mean(), 2))
    neutron_table.cell(2, 0).text = 'Нейтроны, (Nn)/соб.'
    neutron_table.cell(2, 1).text = str(round(neut_stat_info_1.iloc[0].sum(), 2))
    neutron_table.cell(2, 2).text = str(round(neut_stat_info_2.iloc[0].sum(), 2))

    make_table_bold(neutron_table, cols=3, rows=3)
    change_cell_size(neutron_table, column_num=3, size_arr=[2.5, 0.5, 0.5])

    doc.add_paragraph()

    notes = doc.add_paragraph('')
    notes.add_run('Примечание:').bold = True
    doc.add_paragraph(
        '        В таблице 4 представлена сводная информация о неисправностях в работе детекторов кластера.')
    table_title = doc.add_paragraph('Таблица 4: Неисправности.', style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    notes_table = doc.add_table(3, 5, doc.styles['Table Grid'])
    notes_table.cell(0, 0).text = '№'
    notes_table.cell(0, 1).text = 'Кластер'
    notes_table.cell(0, 2).text = '№ Детектора'
    notes_table.cell(0, 3).text = 'Период'
    notes_table.cell(0, 4).text = 'Примечание'
    notes_table.cell(1, 0).text = '1'
    notes_table.cell(2, 0).text = '2'

    change_cell_size(notes_table, column_num=5, size_arr=[0.3, 0.8, 1.2, 1, 4.2])

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    graphic_header = doc.add_paragraph('Продолжительность работы кластеров установки ПРИЗМА-32.', style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    adding_graphic(doc, title='Рис. 1 - Продолжительность работы 1-го кластера в сутки', width=6,
                   picture_path=worktime_pic_path_1)
    adding_graphic(doc, title='Рис. 2 - Продолжительность работы 2-го кластера в сутки', width=6,
                   picture_path=worktime_pic_path_2)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    graphic_header = doc.add_paragraph('Скорость счета событий установки ПРИЗМА-32.', style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    adding_graphic(doc, title='Рис. 3 - Скорость счета событий Fr ≥ 4, A ≥ 5', width=7,
                   picture_path=event_counter_fr_4_pic_path)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    adding_graphic(doc, title='Рис. 4 - Число импульсов в событии, отобранных как нейтрон, при самозапуске кластер 1',
                   width=6, picture_path=n_vs_zero_tr_pic_path_1)
    adding_graphic(doc, title='Рис. 5 - Число импульсов в событии, отобранных как нейтрон, при самозапуске кластер 2',
                   width=6, picture_path=n_vs_zero_tr_pic_path_2)

    table_title = doc.add_paragraph(
        'Таблица 5: Среднее число нейтронов (Nn) для детекторов установки ПРИЗМА-32 за месяц работы, нормированное на количество событий (Ns).(при самозапуске), ',
        style='PItalic')
    table_title.add_run('(100/соб)').bold = True
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    neutron_stat_table = doc.add_table(neut_stat_info_1.shape[0] + neut_stat_info_2.shape[0] + 2,
                                       neut_stat_info_1.shape[1] + 2,
                                       doc.styles['Table Grid'])

    draw_stat_table(neutron_stat_table, neut_stat_info_1, neut_stat_info_2)
    make_stat_table_bold(neutron_stat_table, cols=18, rows=6)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    # Переделать, чтобы амплитуда и кратность были переменными.
    adding_graphic(doc, title='Рис. 6 - Скорость счета  детекторов в 1-м кластере Fr ≥ 2, A > 5',
                   width=6, picture_path=count_rate_amp_5_fr_2_pic_path_1)
    adding_graphic(doc, title='Рис. 7 - Скорость счета  детекторов в 2-м кластере Fr ≥ 2, A > 5',
                   width=6, picture_path=count_rate_amp_5_fr_2_pic_path_2)

    table_title = doc.add_paragraph(
        'Таблица 6: Среднемесячные число срабатываний детекторов установки ПРИЗМА-32, cоб./час.',
        style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    count_rate_stat_info_1, count_rate_stat_info_2 = statistical_table(count_rate_amp_5_fr_2_1, count_rate_amp_5_fr_2_2,
                                                                       dimension='cоб./ч.')

    count_stat_table = doc.add_table(count_rate_stat_info_1.shape[0] + count_rate_stat_info_2.shape[0] + 2,
                                     count_rate_stat_info_1.shape[1] + 2, doc.styles['Table Grid'])

    draw_stat_table(count_stat_table, count_rate_stat_info_1, count_rate_stat_info_2)
    make_stat_table_bold(count_stat_table, cols=18, rows=6)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    adding_graphic(doc, title='Рис. 8 - Скорость счета  детекторов в 1-м кластере Fr ≥ 1, A > 10',
                   width=6, picture_path=count_rate_amp_10_fr_1_pic_path_1)
    adding_graphic(doc, title='Рис. 9 - Скорость счета  детекторов в 2-м кластере Fr ≥ 1, A > 10',
                   width=6, picture_path=count_rate_amp_10_fr_1_pic_path_2)

    table_title = doc.add_paragraph(
        'Таблица 7: Среднемесячные число срабатываний детекторов установки ПРИЗМА-32, cоб./час.',
        style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    count_rate_stat_info_1, count_rate_stat_info_2 = statistical_table(count_rate_amp_10_fr_1_1,
                                                                       count_rate_amp_10_fr_1_2,
                                                                       dimension='cоб./ч.')
    count_stat_table_2 = doc.add_table(count_rate_stat_info_1.shape[0] + count_rate_stat_info_2.shape[0] + 2,
                                       count_rate_stat_info_1.shape[1] + 2, doc.styles['Table Grid'])

    draw_stat_table(count_stat_table_2, count_rate_stat_info_1, count_rate_stat_info_2)
    make_stat_table_bold(count_stat_table_2, cols=18, rows=6)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    graphic_header = doc.add_paragraph(
        'На рисунке 8, 9 представлено число сигналов с А>5 кодов АЦП в час для 16 детекторов.',
        style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    adding_graphic(doc, title='Рис. 10 - Амплитудное распределение сигналов от детекторов, кластер 1 (Fr ≥ 2 и А > 5)',
                   width=6, picture_path=amp_distribution_pic_path_1)
    adding_graphic(doc, title='Рис. 11 - Амплитудное распределение сигналов от детекторов, кластер 2 (Fr ≥ 2 и А > 5)',
                   width=6, picture_path=amp_distribution_pic_path_2)
    add_page_number(doc.sections[0].footer.paragraphs[0])

    doc.save(
        f'{report_path}\\{start_date.day:02}.{start_date.month:02}.{start_date.year}-{end_date.day:02}.{end_date.month:02}.{end_date.year}.docx')

    print(time.time() - t1)
